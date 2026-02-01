#!/usr/bin/env python3
"""
Demo: Incremental Updates with ChromaDB

This script demonstrates how the incremental update feature works:
1. First load: Encodes all Q&A pairs
2. Reload same file: Skips encoding (uses cache)
3. Add new questions: Only encodes the new ones
"""

import os
import sys
from docx import Document


def create_initial_qa_file():
    """Create initial Q&A file with 5 questions"""
    print("Creating initial Q&A file with 5 questions...")
    
    qa_content = [
        ("ما هي ساعات العمل؟", "نحن نعمل من الأحد إلى الخميس من الساعة 9 صباحاً حتى 5 مساءً."),
        ("كيف يمكنني تتبع طلبي؟", "يمكنك تتبع طلبك من خلال الدخول إلى حسابك على الموقع."),
        ("ما هي سياسة الإرجاع؟", "يمكنك إرجاع المنتجات خلال 14 يوماً من تاريخ الاستلام."),
        ("هل تقدمون خدمة التوصيل المجاني؟", "نعم، نقدم توصيل مجاني للطلبات فوق 200 ريال."),
        ("كيف يمكنني تغيير عنوان الشحن؟", "يمكنك تغيير عنوان الشحن من خلال حسابك قبل الشحن."),
    ]
    
    doc = Document()
    doc.add_heading('قاعدة معرفة مركز الاتصال', 0)
    
    for q, a in qa_content:
        doc.add_paragraph(f'سؤال: {q}')
        doc.add_paragraph(f'جواب: {a}')
        doc.add_paragraph()  # Empty line
    
    doc.save('demo_qa.docx')
    print(f"✓ Created demo_qa.docx with {len(qa_content)} Q&A pairs\n")


def add_more_questions():
    """Add 3 more questions to the file"""
    print("Adding 3 new questions to the file...")
    
    doc = Document('demo_qa.docx')
    
    new_qa = [
        ("ما هي طرق الدفع المتاحة؟", "نقبل الدفع عن طريق: البطاقات الائتمانية، مدى، أبل باي."),
        ("هل يمكنني إلغاء طلبي؟", "نعم، يمكنك إلغاء الطلب قبل شحنه من خلال حسابك."),
        ("كم تستغرق مدة الشحن؟", "التوصيل داخل الرياض يستغرق 1-2 يوم عمل."),
    ]
    
    for q, a in new_qa:
        doc.add_paragraph(f'سؤال: {q}')
        doc.add_paragraph(f'جواب: {a}')
        doc.add_paragraph()
    
    doc.save('demo_qa.docx')
    print(f"✓ Added {len(new_qa)} new Q&A pairs to demo_qa.docx\n")


def demo_incremental_updates():
    """Demonstrate incremental updates"""
    from arabic_rag_chromadb import ChromaRAG
    
    # Clean up previous demo database
    import shutil
    if os.path.exists('./demo_chroma_db'):
        shutil.rmtree('./demo_chroma_db')
    
    print("=" * 70)
    print("DEMO: Incremental Updates with ChromaDB")
    print("=" * 70)
    print()
    
    # Step 1: Create initial file
    create_initial_qa_file()
    
    # Step 2: First load - encodes all 5 questions
    print("STEP 1: First Load - Encoding all questions")
    print("-" * 70)
    rag = ChromaRAG(db_path='./demo_chroma_db', collection_name='demo_qa')
    rag.load_from_file('demo_qa.docx')
    print()
    
    # Step 3: Reload same file - should skip encoding
    print("STEP 2: Reload Same File - Should use cache")
    print("-" * 70)
    rag.load_from_file('demo_qa.docx')
    print()
    
    # Step 4: Add new questions
    add_more_questions()
    
    # Step 5: Reload - should only encode the 3 new ones
    print("STEP 3: Reload After Adding Questions - Only encode new ones")
    print("-" * 70)
    rag.load_from_file('demo_qa.docx')
    print()
    
    # Step 6: Test retrieval
    print("STEP 4: Test Retrieval")
    print("-" * 70)
    query = "كيف أدفع؟"
    results = rag.retrieve(query, top_k=2)
    print(f"Query: {query}")
    print(f"Found {len(results)} relevant Q&A pairs:\n")
    for i, (qa, score) in enumerate(results, 1):
        print(f"{i}. Score: {score:.3f}")
        print(f"   Q: {qa['question']}")
        print(f"   A: {qa['answer'][:60]}...")
        print()
    
    # Step 7: Show stats
    print("STEP 5: Database Statistics")
    print("-" * 70)
    stats = rag.get_stats()
    print(f"Total Q&A pairs: {stats['total_qa_pairs']}")
    print(f"Documents tracked: {stats['documents_tracked']}")
    print(f"Last updated: {stats['last_updated']}")
    print(f"Database path: {stats['db_path']}")
    print()
    
    print("=" * 70)
    print("DEMO COMPLETE!")
    print("=" * 70)
    print()
    print("Key Takeaways:")
    print("✓ First load: All 5 Q&A pairs were encoded")
    print("✓ Second load: Skipped encoding (file unchanged)")
    print("✓ After adding 3 new: Only encoded the 3 new pairs")
    print("✓ Total in database: 8 Q&A pairs")
    print()
    print("Files created:")
    print("  - demo_qa.docx (your Q&A file)")
    print("  - demo_chroma_db/ (persistent vector database)")
    print()


def test_with_api():
    """Test with actual Gemini API if available"""
    api_key = os.getenv('GEMINI_API_KEY')
    
    if not api_key:
        print("\n" + "=" * 70)
        print("Skipping API test (GEMINI_API_KEY not set)")
        print("=" * 70)
        return
    
    print("\n" + "=" * 70)
    print("BONUS: Testing with Gemini API")
    print("=" * 70)
    print()
    
    from arabic_rag_chromadb import ArabicCallCenterAgent
    
    agent = ArabicCallCenterAgent(api_key, db_path='./demo_chroma_db')
    agent.load_knowledge_base('demo_qa.docx')
    
    queries = [
        "متى تعملون؟",
        "كيف أدفع؟",
        "هل التوصيل مجاني؟"
    ]
    
    for query in queries:
        print(f"العميل: {query}")
        response = agent.get_response(query)
        print(f"الموظف: {response}")
        print("-" * 70)
        print()


if __name__ == "__main__":
    print("\n🚀 ChromaDB Incremental Updates Demo\n")
    
    try:
        # Run main demo
        demo_incremental_updates()
        
        # Test with API if available
        test_with_api()
        
    except KeyboardInterrupt:
        print("\n\nDemo interrupted by user")
    except Exception as e:
        print(f"\nError: {e}")
        import traceback
        traceback.print_exc()
